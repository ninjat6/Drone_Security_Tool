"""
報告產生器模組

提供低耦合、可維護的 Word 報告產生功能。
採用策略模式與建構者模式設計。
"""

import os
from datetime import datetime
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from constants import (
    TARGETS,
    TARGET_UAV,
    TARGET_GCS,
    PHOTO_ANGLES_ORDER,
    PHOTO_ANGLES_NAME,
    STATUS_PASS,
    STATUS_FAIL,
    STATUS_NA,
    STATUS_UNCHECKED,
    DATE_FMT_PY_FILENAME_SHORT,
)


class ReportDataCollector:
    """
    報告資料收集器

    負責從 ProjectManager 收集所有報告所需資料，
    轉換為結構化格式供 WordDocumentBuilder 使用。
    """

    def __init__(self, pm, config: dict):
        """
        初始化資料收集器

        Args:
            pm: ProjectManager 實例
            config: 規範設定字典
        """
        self.pm = pm
        self.config = config
        self.project_data = pm.project_data
        self.project_path = pm.current_project_path

    def collect_summary(self) -> Dict[str, Dict[str, int]]:
        """
        收集檢測判定摘要（按 UAV/GCS 分類統計）

        即使是同一個檢測項目，若 targets 包含 UAV 和 GCS，
        會分別統計到「無人機檢測」和「地面控制站檢測」中。

        Returns:
            {"無人機檢測共 xx 項": {"pass": n, "fail": n, "na": n, "total": n}, ...}
        """
        # 初始化 UAV 和 GCS 的統計
        target_counts = {
            TARGET_UAV: {"pass": 0, "fail": 0, "na": 0, "total": 0},
            TARGET_GCS: {"pass": 0, "fail": 0, "na": 0, "total": 0},
        }

        for section in self.config.get("test_standards", []):
            sec_id = section["section_id"]

            if not self.pm.is_section_visible(sec_id):
                continue

            for item in section.get("items", []):
                item_uid = item.get("uid", item.get("id"))
                if not self.pm.is_item_visible(item_uid):
                    continue

                targets = item.get("targets", [TARGET_GCS])
                status_map = self.pm.get_test_status_detail(item)

                # 對每個 target 分別統計
                for target in targets:
                    if target not in target_counts:
                        continue

                    target_counts[target]["total"] += 1

                    # 取得該 target 的狀態
                    status = status_map.get(target, "未檢測")

                    if status == "Pass":
                        target_counts[target]["pass"] += 1
                    elif status == "Fail":
                        target_counts[target]["fail"] += 1
                    elif status == "N/A":
                        target_counts[target]["na"] += 1

        # 轉換為報告格式
        summary = {}
        target_names = {
            TARGET_UAV: "無人機檢測",
            TARGET_GCS: "地面控制站檢測",
        }

        for target, counts in target_counts.items():
            if counts["total"] > 0:
                display_name = target_names.get(target, target)
                summary[display_name] = counts

        return summary

    def collect_device_info(self) -> List[Dict[str, str]]:
        """
        動態收集設備資訊（根據 project_meta_schema）

        Returns:
            [{"label": "欄位名稱", "value": "值", "type": "欄位類型"}]
        """
        info_data = self.project_data.get("info", {})
        schema = self.config.get("project_meta_schema", [])

        result = []
        for field in schema:
            field_type = field.get("type", "text")
            # 跳過隱藏欄位和路徑選擇器
            if field_type in ("hidden", "path_selector"):
                continue

            key = field["key"]
            label = field["label"]
            value = info_data.get(key, "")

            # 處理不同類型的值
            if isinstance(value, list):
                value = ", ".join(str(v) for v in value)
            elif value is None:
                value = ""
            else:
                value = str(value)

            result.append(
                {
                    "label": label,
                    "value": value if value else "N/A",
                    "type": field_type,
                }
            )

        return result

    def collect_photos(self) -> Dict[str, Dict[str, Optional[str]]]:
        """
        收集檢測件照片路徑

        Returns:
            {target: {angle: full_path or None}}
        """
        info_data = self.project_data.get("info", {})
        photos = {}

        for target in TARGETS:
            photos[target] = {}
            for angle in PHOTO_ANGLES_ORDER:
                path_key = f"{target}_{angle}_path"
                rel_path = info_data.get(path_key)

                if rel_path and self.project_path:
                    full_path = os.path.join(self.project_path, rel_path)
                    if os.path.exists(full_path):
                        photos[target][angle] = full_path
                    else:
                        photos[target][angle] = None
                else:
                    photos[target][angle] = None

        return photos

    def collect_test_results(self) -> List[Dict[str, Any]]:
        """
        收集各檢測項目結果

        Returns:
            [{
                "id": "6.2.1",
                "name": "身分鑑別",
                "section_name": "系統檢測",
                "targets": ["GCS"],
                "status_map": {"GCS": "Pass"},
                "narrative": {...},
                "result_data": {...}
            }]
        """
        results = []

        for section in self.config.get("test_standards", []):
            sec_id = section["section_id"]
            sec_name = section["section_name"]

            if not self.pm.is_section_visible(sec_id):
                continue

            for item in section.get("items", []):
                item_uid = item.get("uid", item.get("id"))
                item_id = item.get("id", "")

                if not self.pm.is_item_visible(item_uid):
                    continue

                targets = item.get("targets", [TARGET_GCS])
                status_map = self.pm.get_test_status_detail(item)
                narrative = item.get("narrative", {})

                # 收集各 target 的結果資料
                result_data = {}
                meta = self.pm.get_test_meta(item_uid)
                is_shared = meta.get("is_shared", False)

                if is_shared and len(targets) > 1:
                    result_data["Shared"] = self.pm.get_test_result(
                        item_uid, "Shared", True
                    )
                else:
                    for t in targets:
                        result_data[t] = self.pm.get_test_result(item_uid, t, False)

                results.append(
                    {
                        "id": item_id,
                        "name": item.get("name", ""),
                        "section_name": sec_name,
                        "targets": targets,
                        "status_map": status_map,
                        "narrative": narrative,
                        "result_data": result_data,
                        "is_shared": is_shared,
                    }
                )

        return results

    def collect_all(self) -> Dict[str, Any]:
        """
        整合所有報告資料

        Returns:
            完整的報告資料字典
        """
        project_name = self.project_data.get("info", {}).get(
            "project_name", "未命名專案"
        )

        return {
            "project_name": project_name,
            "project_path": self.project_path,
            "standard_name": self.config.get("standard_name", ""),
            "standard_version": self.config.get("standard_version", ""),
            "generate_time": datetime.now(),
            "summary": self.collect_summary(),
            "device_info": self.collect_device_info(),
            "photos": self.collect_photos(),
            "test_results": self.collect_test_results(),
        }


class WordDocumentBuilder:
    """
    Word 文件建構器

    接收結構化資料，建構 Word 文件。
    """

    def __init__(self, data: Dict[str, Any]):
        """
        初始化建構器

        Args:
            data: 由 ReportDataCollector.collect_all() 產生的資料
        """
        self.data = data
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """設定文件基本格式"""
        # 設定預設字型（處理中文）
        self.doc.styles["Normal"].font.name = "標楷體"
        self.doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        self.doc.styles["Normal"].font.size = Pt(12)

    def _set_cell_shading(self, cell, color: str):
        """設定儲存格背景色"""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _add_heading(self, text: str, level: int = 1):
        """新增標題"""
        heading = self.doc.add_heading(text, level=level)
        # 設定中文字型
        for run in heading.runs:
            run.font.name = "標楷體"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    def build_summary_section(self):
        """建構檢測判定摘要區塊"""
        self._add_heading("檢測判定摘要", level=1)

        summary = self.data.get("summary", {})
        if not summary:
            self.doc.add_paragraph("無檢測資料")
            return

        # 建立摘要表格
        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 標題列
        header_cells = table.rows[0].cells
        header_cells[0].text = "檢測類型"
        header_cells[1].text = "檢測判定"

        for cell in header_cells:
            self._set_cell_shading(cell, "4472C4")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.color.rgb = None  # 白色
                    run.bold = True

        # 資料列
        for sec_name, counts in summary.items():
            row = table.add_row()
            row.cells[0].text = f"{sec_name}共 {counts['total']} 項"

            # 判定結果
            result_text = (
                f"通過基準要求: {counts['pass']} 項\n"
                f"不通過基準要求: {counts['fail']} 項\n"
                f"不適用: {counts['na']} 項"
            )
            row.cells[1].text = result_text

        self.doc.add_paragraph()  # 空行

    def build_device_info_section(self):
        """建構檢測件資料區塊"""
        self._add_heading("檢測件資料", level=1)

        device_info = self.data.get("device_info", [])
        if not device_info:
            self.doc.add_paragraph("無設備資料")
            return

        # 建立資料表格
        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 標題列
        header_cells = table.rows[0].cells
        header_cells[0].text = "項目"
        header_cells[1].text = "資訊"

        for cell in header_cells:
            self._set_cell_shading(cell, "4472C4")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        # 資料列
        for item in device_info:
            row = table.add_row()
            row.cells[0].text = item["label"]
            row.cells[1].text = item["value"]
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # 空行

    def build_photo_section(self):
        """建構檢測件照片區塊"""
        self._add_heading("檢測件照片", level=1)

        photos = self.data.get("photos", {})

        for target in TARGETS:
            target_name = "遙控無人機" if target == TARGET_UAV else "地面控制站"
            self._add_heading(f"檢測件{target_name}", level=2)

            target_photos = photos.get(target, {})
            photo_count = 1

            for angle in PHOTO_ANGLES_ORDER:
                angle_name = PHOTO_ANGLES_NAME.get(angle, angle)
                photo_path = target_photos.get(angle)

                if photo_path and os.path.exists(photo_path):
                    # 插入照片
                    try:
                        self.doc.add_picture(photo_path, width=Inches(4))
                        last_paragraph = self.doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        self.doc.add_paragraph(f"[無法載入圖片: {angle_name}]")
                else:
                    # 顯示未上傳
                    p = self.doc.add_paragraph(f"[未上傳]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 圖片標題
                caption = self.doc.add_paragraph(
                    f"圖 {photo_count} {target_name} - {angle_name}"
                )
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                photo_count += 1

            self.doc.add_paragraph()  # 空行

    def build_test_results_section(self):
        """建構檢測項目結果區塊（先 UAV 再 GCS）"""
        test_results = self.data.get("test_results", [])
        if not test_results:
            self._add_heading("檢測執行過程", level=1)
            self.doc.add_paragraph("無檢測項目")
            return

        # 定義 target 順序和名稱
        target_order = [
            (TARGET_UAV, "遙控無人機資安檢測執行過程"),
            (TARGET_GCS, "地面控制站資安檢測執行過程"),
        ]

        for target, target_title in target_order:
            # 檢查此 target 是否有任何項目
            target_items = [
                item for item in test_results if target in item.get("targets", [])
            ]

            if not target_items:
                continue

            # 顯示 target 大標題
            self._add_heading(target_title, level=1)

            # 按 section 分組
            current_section = ""
            section_count = 0

            for item in target_items:
                # 區塊標題
                if item["section_name"] != current_section:
                    current_section = item["section_name"]
                    section_count += 1
                    section_title = f"{section_count}. {current_section}"
                    self._add_heading(section_title, level=2)

                # 建立檢測項目表格（傳入目標 target）
                self._build_single_test_table(item, target)
                self.doc.add_paragraph()  # 空行

    def _build_single_test_table(self, item: Dict[str, Any], target: str):
        """建構單一檢測項目的表格（針對特定 target）"""
        # 取得該 target 的判定結果
        status_map = item.get("status_map", {})
        target_status = status_map.get(target, "未檢測")

        # 轉換狀態文字
        if target_status == "Pass":
            result_text = "通過"
        elif target_status == "Fail":
            result_text = "不通過"
        elif target_status == "N/A":
            result_text = "不適用"
        else:
            result_text = "未檢測"

        narrative = item.get("narrative", {})
        result_data = item.get("result_data", {})
        is_shared = item.get("is_shared", False)

        # 取得該 target 的結果資料（如果是共用模式，使用 Shared 的資料）
        if is_shared:
            target_result_data = result_data.get("Shared", {})
        else:
            target_result_data = result_data.get(target, {})

        # 建立表格
        table = self.doc.add_table(rows=6, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 第一列：編號、檢測項目、判定
        row0 = table.rows[0]
        row0.cells[0].text = "編號"
        row0.cells[1].text = "檢測項目"
        row0.cells[2].text = "判定"
        for cell in row0.cells:
            self._set_cell_shading(cell, "4472C4")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True

        # 第二列：編號值、項目名稱、判定結果
        row1 = table.rows[1]
        row1.cells[0].text = item.get("id", "")
        row1.cells[1].text = item.get("name", "")
        row1.cells[2].text = result_text
        for cell in row1.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 第三列：檢測目的（合併儲存格）
        row2 = table.rows[2]
        row2.cells[0].merge(row2.cells[2])
        row2.cells[0].text = "檢測目的"
        self._set_cell_shading(row2.cells[0], "D9E2F3")
        row2.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 第四列：檢測目的內容
        row3 = table.rows[3]
        row3.cells[0].merge(row3.cells[2])
        row3.cells[0].text = narrative.get("purpose", "")

        # 第五列：檢測方法標題
        row4 = table.rows[4]
        row4.cells[0].merge(row4.cells[2])
        row4.cells[0].text = "檢測方法"
        self._set_cell_shading(row4.cells[0], "D9E2F3")
        row4.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 第六列：檢測方法內容
        row5 = table.rows[5]
        row5.cells[0].merge(row5.cells[2])
        row5.cells[0].text = narrative.get("method", "")

        # 新增判定標準
        criteria_header = table.add_row()
        criteria_header.cells[0].merge(criteria_header.cells[2])
        criteria_header.cells[0].text = "檢測判定"
        self._set_cell_shading(criteria_header.cells[0], "D9E2F3")
        criteria_header.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        criteria_content = table.add_row()
        criteria_content.cells[0].merge(criteria_content.cells[2])
        criteria_content.cells[0].text = narrative.get("criteria", "")

        # 新增過程說明標題
        process_header = table.add_row()
        process_header.cells[0].merge(process_header.cells[2])
        process_header.cells[0].text = "檢測過程說明"
        self._set_cell_shading(process_header.cells[0], "D9E2F3")
        process_header.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 過程說明內容：先放圖片，再放判定原因（都在表格內）
        process_content = table.add_row()
        process_content.cells[0].merge(process_content.cells[2])
        cell = process_content.cells[0]

        # 從該 target 的結果資料收集附件和判定原因
        all_attachments = []
        note_text = ""

        if isinstance(target_result_data, dict):
            # 收集附件
            attachments = target_result_data.get("attachments", [])
            all_attachments.extend(attachments)

            # 收集判定原因
            note_text = target_result_data.get(
                "description", target_result_data.get("note", "")
            )

        # 先在表格內插入圖片
        img_count = 1
        for att in all_attachments:
            att_path = att.get("path", "")
            att_title = att.get("title", "")
            att_type = att.get("type", "image")

            # 如果是相對路徑，轉換為絕對路徑
            if not os.path.isabs(att_path) and self.data.get("project_path"):
                att_path = os.path.join(self.data["project_path"], att_path)

            # 只處理圖片類型
            if att_type == "image" and os.path.exists(att_path):
                try:
                    # 在表格格子內插入圖片
                    paragraph = (
                        cell.paragraphs[0] if img_count == 1 else cell.add_paragraph()
                    )
                    run = paragraph.add_run()
                    run.add_picture(att_path, width=Inches(5))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    paragraph = cell.add_paragraph(f"[無法載入圖片]")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 圖片標題
                caption_text = (
                    f"圖 {img_count} {att_title}" if att_title else f"圖 {img_count}"
                )
                caption_para = cell.add_paragraph(caption_text)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_count += 1

        # 再放判定原因文字
        if note_text:
            # 如果已經有圖片，新增一個段落放文字；否則直接設定格子文字
            if img_count > 1:
                note_para = cell.add_paragraph(note_text)
            else:
                cell.text = note_text

        # 新增檢測人員和檢測工具名稱（使用新的表格來確保各佔一半）
        # 結束目前表格，新增一個 2 欄的表格
        footer_table = self.doc.add_table(rows=2, cols=2)
        footer_table.style = "Table Grid"
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 設定欄寬各佔一半
        for row in footer_table.rows:
            for cell in row.cells:
                cell.width = Cm(8)  # 各約 8cm，總共 16cm

        # 標題列
        footer_table.rows[0].cells[0].text = "檢測人員"
        footer_table.rows[0].cells[0].paragraphs[
            0
        ].alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_table.rows[0].cells[1].text = "檢測工具名稱"
        footer_table.rows[0].cells[1].paragraphs[
            0
        ].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_cell_shading(footer_table.rows[0].cells[0], "D9E2F3")
        self._set_cell_shading(footer_table.rows[0].cells[1], "D9E2F3")

        # 取得檢測人員名稱（從專案資料的 tester_1）
        tester_name = self.data.get("device_info", [])
        tester_1 = ""
        for info in tester_name:
            if info.get("label") == "檢測人員1":
                tester_1 = info.get("value", "")
                break

        # 值列
        footer_table.rows[1].cells[0].text = (
            tester_1 if tester_1 and tester_1 != "N/A" else ""
        )
        footer_table.rows[1].cells[0].paragraphs[
            0
        ].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 檢測工具名稱：從 target_result_data 取得，如果判定為不適用則顯示 N/A
        tool_name = ""
        if isinstance(target_result_data, dict):
            tool_name = target_result_data.get("tool_name", "")

        if result_text == "不適用":
            tool_name = "N/A"

        footer_table.rows[1].cells[1].text = tool_name if tool_name else ""
        footer_table.rows[1].cells[1].paragraphs[
            0
        ].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def build(self) -> Document:
        """
        建構完整文件

        Returns:
            完整的 Document 物件
        """
        self.build_summary_section()
        self.build_device_info_section()
        self.build_photo_section()
        self.build_test_results_section()
        return self.doc

    def save(self, output_path: str) -> str:
        """
        儲存文件

        Args:
            output_path: 輸出路徑

        Returns:
            實際儲存的路徑
        """
        self.build()
        self.doc.save(output_path)
        return output_path


def generate_report(pm, config: dict) -> str:
    """
    產生檢測報告（高層 API）

    Args:
        pm: ProjectManager 實例
        config: 規範設定字典

    Returns:
        報告檔案路徑

    Raises:
        ValueError: 如果專案未開啟
    """
    if not pm.current_project_path:
        raise ValueError("請先開啟專案")

    # 收集資料
    collector = ReportDataCollector(pm, config)
    data = collector.collect_all()

    # 建構文件
    builder = WordDocumentBuilder(data)

    # 產生檔名
    project_name = data.get("project_name", "報告")
    date_str = datetime.now().strftime(DATE_FMT_PY_FILENAME_SHORT)
    filename = f"{project_name}_檢測報告_{date_str}.docx"

    # 清理檔名中的特殊字元
    for char in ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]:
        filename = filename.replace(char, "_")

    output_path = os.path.join(pm.current_project_path, filename)

    # 儲存
    return builder.save(output_path)
